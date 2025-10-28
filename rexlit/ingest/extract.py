"""Document content extraction for various file formats."""

from pathlib import Path

from pydantic import BaseModel, Field


class ExtractedContent(BaseModel):
    """Extracted content from a document."""

    path: str = Field(..., description="Source document path")
    text: str = Field(..., description="Extracted text content")
    page_count: int | None = Field(None, description="Number of pages (if applicable)")
    metadata: dict[str, str] = Field(
        default_factory=dict, description="Additional document metadata"
    )


def extract_text_file(file_path: Path) -> ExtractedContent:
    """Extract text from plain text file.

    Args:
        file_path: Path to text file

    Returns:
        Extracted content

    Raises:
        FileNotFoundError: If file does not exist
        UnicodeDecodeError: If file cannot be decoded
    """
    with open(file_path, encoding="utf-8", errors="replace") as f:
        text = f.read()

    return ExtractedContent(
        path=str(file_path.absolute()),
        text=text,
        metadata={"format": "text"},
    )


def extract_pdf(file_path: Path) -> ExtractedContent:
    """Extract text from PDF file.

    Args:
        file_path: Path to PDF file

    Returns:
        Extracted content

    Raises:
        FileNotFoundError: If file does not exist
        ImportError: If PyMuPDF is not installed
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise ImportError(
            "PyMuPDF is required for PDF extraction. " "Install with: pip install pymupdf"
        ) from e

    doc = fitz.open(file_path)
    text_parts = []
    metadata_dict = {}

    # Extract text from each page
    for page_num in range(len(doc)):
        page = doc[page_num]
        text_parts.append(page.get_text())

    # Extract PDF metadata
    pdf_metadata = doc.metadata
    if pdf_metadata:
        for key, value in pdf_metadata.items():
            if value:
                metadata_dict[key.lower()] = str(value)

    doc.close()

    return ExtractedContent(
        path=str(file_path.absolute()),
        text="\n\n".join(text_parts),
        page_count=len(doc),
        metadata={"format": "pdf", **metadata_dict},
    )


def extract_docx(file_path: Path) -> ExtractedContent:
    """Extract text from DOCX file.

    Args:
        file_path: Path to DOCX file

    Returns:
        Extracted content

    Raises:
        FileNotFoundError: If file does not exist
        ImportError: If python-docx is not installed
    """
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX extraction. " "Install with: pip install python-docx"
        ) from e

    doc = Document(file_path)
    text_parts = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)

    # Extract core properties metadata
    metadata_dict = {}
    core_props = doc.core_properties
    if core_props.author:
        metadata_dict["author"] = core_props.author
    if core_props.title:
        metadata_dict["title"] = core_props.title
    if core_props.subject:
        metadata_dict["subject"] = core_props.subject
    if core_props.created:
        metadata_dict["created"] = core_props.created.isoformat()
    if core_props.modified:
        metadata_dict["modified"] = core_props.modified.isoformat()

    return ExtractedContent(
        path=str(file_path.absolute()),
        text="\n\n".join(text_parts),
        metadata={"format": "docx", **metadata_dict},
    )


def extract_image(file_path: Path) -> ExtractedContent:
    """Extract metadata from image file (no text extraction).

    Args:
        file_path: Path to image file

    Returns:
        Extracted content (empty text, image metadata only)

    Raises:
        FileNotFoundError: If file does not exist
        ImportError: If Pillow is not installed
    """
    try:
        from PIL import Image
    except ImportError as e:
        raise ImportError(
            "Pillow is required for image handling. " "Install with: pip install pillow"
        ) from e

    img = Image.open(file_path)
    metadata_dict = {
        "format": img.format or "unknown",
        "mode": img.mode,
        "size": f"{img.width}x{img.height}",
    }

    # Extract EXIF data if available
    if hasattr(img, "_getexif") and img._getexif():
        exif = img._getexif()
        if exif:
            metadata_dict["has_exif"] = "true"

    img.close()

    return ExtractedContent(
        path=str(file_path.absolute()),
        text="",  # Images require OCR for text extraction
        metadata=metadata_dict,
    )


def extract_document(file_path: Path) -> ExtractedContent:
    """Extract content from document based on file type.

    Args:
        file_path: Path to document

    Returns:
        Extracted content

    Raises:
        FileNotFoundError: If file does not exist
        ValueError: If file format is not supported
    """
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    extension = file_path.suffix.lower()

    # Route to appropriate extractor
    if extension == ".pdf":
        return extract_pdf(file_path)
    elif extension == ".docx":
        return extract_docx(file_path)
    elif extension in {".txt", ".md", ".log", ".csv"}:
        return extract_text_file(file_path)
    elif extension in {".png", ".jpg", ".jpeg", ".tiff", ".bmp", ".gif"}:
        return extract_image(file_path)
    else:
        raise ValueError(f"Unsupported file format: {extension}")
